from __future__ import annotations
from pathlib import Path
import pandas as pd
from pypdf import PdfReader
from docx import Document


def parse_pdf(path: str | Path) -> str:
    reader = PdfReader(str(path))
    pages = []
    for page in reader.pages:
        pages.append(page.extract_text() or "")
    return "\n".join(pages)


def parse_docx(path: str | Path) -> str:
    doc = Document(str(path))
    parts = [p.text for p in doc.paragraphs if p.text]
    for table in doc.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(parts)


def parse_xlsx(path: str | Path) -> str:
    xl = pd.ExcelFile(path)
    parts = []
    for sheet in xl.sheet_names:
        df = pd.read_excel(path, sheet_name=sheet, header=None)
        parts.append(f"# Лист: {sheet}")
        parts.append(df.fillna("").astype(str).to_csv(sep="\t", index=False, header=False))
    return "\n".join(parts)


def parse_text_file(path: str | Path) -> str:
    return Path(path).read_text(encoding="utf-8", errors="ignore")


def parse_tender_file(path: str | Path) -> str:
    path = Path(path)
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return parse_pdf(path)
    if suffix == ".docx":
        return parse_docx(path)
    if suffix in {".xlsx", ".xls"}:
        return parse_xlsx(path)
    if suffix in {".txt", ".md"}:
        return parse_text_file(path)
    raise ValueError(f"Unsupported file type: {suffix}. Use PDF, DOCX, XLSX, TXT or MD.")
